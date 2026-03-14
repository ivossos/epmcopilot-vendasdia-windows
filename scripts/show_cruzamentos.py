#!/usr/bin/env python3
"""
Mostra os cruzamentos completos onde há dados no dashboard_data.json.
Gera relatório em Word (data/cruzamentos.docx).
Run: python scripts/show_cruzamentos.py
"""
import json
from pathlib import Path
from collections import defaultdict

PROJECT = Path(__file__).resolve().parent.parent
DATA_PATH = PROJECT / "data" / "dashboard_data.json"
OUT_DOCX = PROJECT / "data" / "cruzamentos.docx"


def main():
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Instale python-docx: pip install python-docx")
        return 1

    if not DATA_PATH.exists():
        print(f"Arquivo não encontrado: {DATA_PATH}")
        print("Execute primeiro: python scripts/fetch_dashboard_data.py")
        return 1

    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    records = data.get("records", [])

    if not records:
        print("Nenhum registro encontrado.")
        return 1

    # Dimensões únicas
    scenarios = sorted(set(r.get("scenario", "Real/Trabalho") for r in records), key=lambda s: (s != "Real/Trabalho", s))
    years = sorted(set(r["year"] for r in records))
    months = sorted(set(r["month"] for r in records), key=lambda m: ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"].index(m) if m in ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"] else 99)
    filiais = sorted(set(r["filial"] for r in records), key=lambda f: (f != "All BU", f))
    categorias = sorted(set(r.get("categoria", "—") for r in records), key=lambda c: (c != "Total Categoria", c != "All Categoria", c))
    accounts = sorted(set(r["account"] for r in records))

    # Cruzamentos
    by_cross = defaultdict(lambda: defaultdict(list))
    for r in records:
        key = (r.get("scenario", "Real/Trabalho"), r["year"], r["month"], r["filial"], r.get("categoria", "—"))
        by_cross[key][r["account"]].append(r["value"])

    by_filial = defaultdict(int)
    by_scenario = defaultdict(int)
    for r in records:
        by_filial[r["filial"]] += 1
        by_scenario[r.get("scenario", "Real/Trabalho")] += 1

    by_cat = defaultdict(int)
    for r in records:
        by_cat[r.get("categoria", "—")] += 1

    # Gerar Word
    doc = Document()
    doc.add_heading("Cruzamentos Completos Onde Há Dados", 0)

    p = doc.add_paragraph()
    p.add_run(f"Fonte: {data.get('source', 'EPBCS VendaDia')} | Cenários: {', '.join(data.get('scenarios', scenarios))}\n")
    p.add_run(f"Total de registros: {len(records):,}")

    doc.add_heading("Dimensões Encontradas", level=1)
    doc.add_paragraph(f"Cenários: {', '.join(scenarios)}")
    doc.add_paragraph(f"Anos: {', '.join(years)}")
    doc.add_paragraph(f"Meses: {', '.join(months)}")
    doc.add_paragraph(f"Filiais: {', '.join(filiais)}")
    doc.add_paragraph(f"Categorias: {', '.join(categorias)}")
    doc.add_paragraph(f"Contas: {', '.join(accounts)}")

    doc.add_heading("Cruzamentos (Cenário | Ano | Mês | Filial | Categoria | Contas com dados)", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Cenário"
    hdr[1].text = "Ano"
    hdr[2].text = "Mês"
    hdr[3].text = "Filial"
    hdr[4].text = "Categoria"
    hdr[5].text = "Contas com dados"

    for (scn, yr, mo, fil, cat), accs in sorted(by_cross.items()):
        row = table.add_row()
        row.cells[0].text = scn
        row.cells[1].text = yr
        row.cells[2].text = mo
        row.cells[3].text = fil
        row.cells[4].text = cat
        row.cells[5].text = ", ".join(sorted(accs.keys()))

    doc.add_heading("Resumo por Cenário", level=1)
    t_scn = doc.add_table(rows=1, cols=2)
    t_scn.style = "Table Grid"
    t_scn.rows[0].cells[0].text = "Cenário"
    t_scn.rows[0].cells[1].text = "Registros"
    for scn in scenarios:
        r = t_scn.add_row()
        r.cells[0].text = scn
        r.cells[1].text = f"{by_scenario[scn]:,}"

    doc.add_heading("Resumo por Filial", level=1)
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Filial"
    t2.rows[0].cells[1].text = "Registros"
    for fil in filiais:
        r = t2.add_row()
        r.cells[0].text = fil
        r.cells[1].text = f"{by_filial[fil]:,}"

    doc.add_heading("Resumo por Categoria", level=1)
    t3 = doc.add_table(rows=1, cols=2)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].text = "Categoria"
    t3.rows[0].cells[1].text = "Registros"
    for cat in categorias:
        r = t3.add_row()
        r.cells[0].text = cat
        r.cells[1].text = f"{by_cat[cat]:,}"

    doc.add_heading("Comandos de Exportação dos Dados", level=1)
    doc.add_paragraph("1. Exportar dados do EPBCS para o dashboard (JSON):")
    p = doc.add_paragraph()
    p.add_run(".venv/bin/python scripts/fetch_dashboard_data.py").bold = True
    doc.add_paragraph("   Saída: data/dashboard_data.json — Dados de vendas (Real + Orçamento, anos, meses, filiais, categorias, contas).")

    doc.add_paragraph("2. Gerar relatório de cruzamentos (este documento):")
    p2 = doc.add_paragraph()
    p2.add_run(".venv/bin/python scripts/show_cruzamentos.py").bold = True
    doc.add_paragraph("   Saída: data/cruzamentos.docx — Relatório com cruzamentos completos onde há dados.")

    doc.add_paragraph("3. Buscar dados no EPBCS (exploratório):")
    p3 = doc.add_paragraph()
    p3.add_run("python scripts/find_data.py").bold = True
    doc.add_paragraph("   Explora plan types VendaDia/Diario, cenários e anos para encontrar intersecções com dados.")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"Relatório salvo em: {OUT_DOCX}")
    return 0


if __name__ == "__main__":
    exit(main())
